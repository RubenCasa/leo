# -*- coding: utf-8 -*-
"""
Generador de Informe Word (.docx) - Proyecto LEO-CONNECT
Lácteos Leo - Plataforma Web E-Commerce
Formato APA 7.ª edición con datos COCOMO REALES del equipo UNACH
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ============================================================
# CONFIGURACIÓN DEL DOCUMENTO
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)
pf = style.paragraph_format
pf.line_spacing = 2.0
pf.space_after = Pt(0)
pf.space_before = Pt(0)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.54)

# ============================================================
# FUNCIONES AUXILIARES
# ============================================================
def add_heading_apa(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs:
            run.font.size = Pt(14)
            run.bold = True
    elif level == 2:
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in h.runs:
            run.font.size = Pt(12)
            run.bold = True
    elif level == 3:
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in h.runs:
            run.font.size = Pt(12)
            run.bold = True
            run.italic = True
    return h

def add_p(text, bold=False, italic=False, align='justify'):
    p = doc.add_paragraph()
    if align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p

def add_mixed(parts):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic
    return p

def add_bullet(text, bold_prefix="", level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.bold = True
        r2 = p.add_run(text)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
    else:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
    return p

def add_numbered(number, bold_part, rest):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Cm(1.27)
    p.add_run(f"{number}. ").font.name = 'Times New Roman'
    rb = p.add_run(bold_part)
    rb.font.name = 'Times New Roman'
    rb.font.size = Pt(12)
    rb.bold = True
    rr = p.add_run(rest)
    rr.font.name = 'Times New Roman'
    rr.font.size = Pt(12)

def set_cell_shade(cell, color):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:fill'), color)
    sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)

def make_table(headers, rows, header_color='1A5276'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shade(cell, header_color)
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(ct))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            if ri % 2 == 1:
                set_cell_shade(cell, 'EBF5FB')
    doc.add_paragraph()
    return table

def make_table_teal(headers, rows):
    """Tabla con encabezado verde-azulado como en la imagen COCOMO"""
    return make_table(headers, rows, header_color='008080')

def add_formula(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = True
    doc.add_paragraph()

def add_placeholder(title, desc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"[{title}]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = True
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    run2 = p2.add_run(desc)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(10)
    run2.italic = True
    run2.font.color.rgb = RGBColor(120, 120, 120)

def add_ref(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.left_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def pb():
    doc.add_page_break()


# ============================================================
#                     PORTADA
# ============================================================
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("UNIVERSIDAD NACIONAL DE CHIMBORAZO")
r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Facultad de Ingeniería\nCarrera de Ingeniería en Sistemas y Computación")
r.font.name = 'Times New Roman'; r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Asignatura: Ingeniería de Software")
r.font.name = 'Times New Roman'; r.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LEO-CONNECT: PLATAFORMA WEB DE E-COMMERCE\nCON FACTURACIÓN ELECTRÓNICA SRI\nPARA LÁCTEOS LEO")
r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Informe de Investigación Formativa – Formato APA 7.ª edición")
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Integrantes del equipo:")
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True

for nombre in ["Rubén Darío Casa Casa", "[Nombre Integrante 2]", "[Nombre Integrante 3]", "[Nombre Integrante 4]"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(nombre)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Docente:")
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("[Nombre del Docente]")
r.font.name = 'Times New Roman'; r.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("5to Semestre – Periodo Académico 2026\nFecha: Julio de 2026")
r.font.name = 'Times New Roman'; r.font.size = Pt(12)

pb()

# ============================================================
# TABLA DE CONTENIDOS
# ============================================================
add_heading_apa("TABLA DE CONTENIDOS", level=1)

toc = [
    ("7.", "DESARROLLO DEL INFORME"),
    ("   7.1", "Introducción"),
    ("   7.2", "Descripción de la metodología"),
    ("   7.3", "Descripción de las acciones realizadas"),
    ("   7.4", "Resultados"),
    ("   7.5", "Bibliografía"),
    ("8.", "ANEXOS (EVIDENCIAS)"),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(f"{num}  {title}")
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

pb()

# ============================================================
# 7. DESARROLLO DEL INFORME
# ============================================================
add_heading_apa("7. DESARROLLO DEL INFORME", level=1)

# ============================================================
# 7.1 INTRODUCCIÓN
# ============================================================
add_heading_apa("7.1 Introducción", level=2)

add_mixed([
    ("El presente informe documenta el proceso de investigación formativa, diseño, desarrollo e implementación de ", False, False),
    ("LEO-CONNECT", True, False),
    (", una plataforma web de comercio electrónico con facturación electrónica SRI, desarrollada para la empresa ", False, False),
    ("Lácteos Leo", True, False),
    (", productora artesanal de quesos, yogures y bebidas naturales con sede en Latacunga, provincia de Cotopaxi, Ecuador, fundada en el año 1990. El proyecto fue desarrollado por un equipo de ", False, False),
    ("4 estudiantes de 5to semestre de la Universidad Nacional de Chimborazo (UNACH)", True, False),
    (" como parte de la asignatura de Ingeniería de Software.", False, False),
])

add_mixed([
    ("El proyecto surge como respuesta a la necesidad de modernizar la presencia digital de la empresa, que hasta el momento contaba únicamente con un sitio web estático de carácter informativo (disponible en ", False, False),
    ("lacteosleo.com", False, True),
    ("), cuya funcionalidad se limitaba a funcionar como un dashboard interactivo para que los clientes conozcan los productos y se contacten con el propietario para concretar ventas de manera manual a través de WhatsApp.", False, False),
])

add_mixed([
    ("Ante este escenario, el equipo planteó la creación de una plataforma digital completa que integrara las funcionalidades de un ", False, False),
    ("e-commerce", False, True),
    (" moderno: catálogo interactivo con carrito de compras, sistema de autenticación multirol, pasarela de pago simulada con protocolo de seguridad 3D Secure 2.0, generación automatizada de comprobantes electrónicos compatibles con el estándar del Servicio de Rentas Internas (SRI) del Ecuador en formato XML con clave de acceso de 49 dígitos, panel de administración para la gestión integral de inventarios, pedidos y usuarios, y un panel de vendedor para atención presencial tipo punto de venta (POS).", False, False),
])

add_p("Para el desarrollo de la plataforma se emplearon las siguientes herramientas y tecnologías:")

tools = [
    ("React 19 con TypeScript: ", "Biblioteca de interfaz de usuario que permite crear componentes reactivos y mantener una arquitectura de código estructurada y escalable."),
    ("Vite 8: ", "Empaquetador y servidor de desarrollo de alto rendimiento que agilizó el proceso de compilación y recarga en caliente durante el desarrollo."),
    ("Tailwind CSS 4 y CSS Vanilla: ", "Frameworks de estilos que facilitaron la construcción de una interfaz visualmente atractiva, responsiva y compatible con diferentes dispositivos."),
    ("Supabase (PostgreSQL + Auth + Row Level Security): ", "Plataforma de backend como servicio que proporcionó la base de datos relacional, el sistema de autenticación segura con OAuth 2.0 (correo/contraseña y Google), y las políticas de seguridad a nivel de fila (RLS) para el control de acceso por roles."),
    ("React Router DOM v7: ", "Librería de enrutamiento del lado del cliente que habilitó la navegación tipo SPA sin recargas de página."),
    ("jsPDF y jspdf-autotable: ", "Bibliotecas de JavaScript para la generación de documentos PDF dinámicos, utilizadas en la emisión de facturas electrónicas en formato RIDE."),
    ("Lucide React: ", "Librería de iconografía vectorial moderna empleada en toda la interfaz de usuario."),
    ("Google Gemini y ChatGPT (IA Generativa): ", "Herramientas de inteligencia artificial utilizadas como asistentes de programación. Según la estimación COCOMO Intermedio, el proyecto requeriría aproximadamente 7 desarrolladores; sin embargo, gracias a la IA el equipo de 4 estudiantes logró compensar esta diferencia de productividad de manera significativa."),
    ("Vercel: ", "Plataforma de despliegue continuo utilizada para la publicación de la aplicación en producción con HTTPS automático y CDN global."),
    ("Git y GitHub: ", "Sistema de control de versiones y repositorio remoto para la gestión colaborativa del código fuente del equipo."),
]
for bp, nt in tools:
    add_bullet(nt, bold_prefix=bp)

add_mixed([
    ("El proyecto fue abordado con una metodología de desarrollo ágil adaptada al contexto académico, combinando investigación exploratoria del estado del arte en ", False, False),
    ("e-commerce", False, True),
    (" para pequeñas y medianas empresas del sector lácteo, con ciclos iterativos de diseño, implementación y prueba. La estimación con el modelo COCOMO Básico e Intermedio confirmó que la complejidad del sistema, especialmente el módulo de facturación electrónica SRI, elevaba el esfuerzo requerido un ", False, False),
    ("+21.64%", True, False),
    (" respecto al cálculo genérico, justificando el uso estratégico de IA generativa como herramienta complementaria de desarrollo.", False, False),
])

pb()

# ============================================================
# 7.2 DESCRIPCIÓN DE LA METODOLOGÍA
# ============================================================
add_heading_apa("7.2 Descripción de la metodología", level=2)

add_heading_apa("7.2.1 Tipo de investigación y enfoque metodológico", level=3)

add_mixed([
    ("La investigación desarrollada es de tipo ", False, False),
    ("aplicada", True, False),
    (" con enfoque ", False, False),
    ("mixto", True, False),
    (": cualitativo en la fase de levantamiento de requisitos y análisis del contexto organizacional de Lácteos Leo; y cuantitativo en la estimación del esfuerzo de desarrollo mediante el modelo COCOMO (Constructive Cost Model) en sus variantes Básica e Intermedia, así como en la medición de métricas de rendimiento de la aplicación.", False, False),
])

add_heading_apa("7.2.2 Modelo de desarrollo de software", level=3)

add_mixed([
    ("Se adoptó un modelo de desarrollo ", False, False),
    ("iterativo-incremental", True, False),
    (", compatible con los principios del marco ágil, que permitió entregar funcionalidades de forma progresiva y validar cada módulo con retroalimentación constante. El proceso se organizó en las siguientes fases:", False, False),
])

fases = [
    ("Fase 1 – Análisis y planificación: ", "Estudio del sitio web existente de Lácteos Leo (lacteosleo.com), entrevistas con el propietario para identificar necesidades y oportunidades de mejora, levantamiento de requisitos funcionales y no funcionales, definición del alcance del sistema y elaboración del plan de proyecto."),
    ("Fase 2 – Diseño del sistema: ", "Elaboración de diagramas UML (casos de uso, clases, secuencia y actividad), diseño de la arquitectura de software (modelo MVC adaptado a componentes React), diseño del esquema de base de datos relacional en PostgreSQL y diseño de las interfaces de usuario (mockups)."),
    ("Fase 3 – Desarrollo e implementación: ", "Codificación de los 6 módulos del sistema en ciclos de sprint semanales: módulo de e-commerce, facturación electrónica SRI, autenticación y seguridad, gestión de inventario, reportes y exportación, e infraestructura core MVC."),
    ("Fase 4 – Pruebas y validación: ", "Ejecución de pruebas funcionales, de usabilidad y de seguridad sobre cada módulo desarrollado."),
    ("Fase 5 – Despliegue y socialización: ", "Publicación de la aplicación en Vercel, capacitación al propietario de Lácteos Leo y presentación de resultados al equipo docente."),
]
for i, (bp, nt) in enumerate(fases, 1):
    add_numbered(i, bp, nt)

# ===== COCOMO =====
add_heading_apa("7.2.3 Estimación con modelo COCOMO", level=3)

add_mixed([
    ("Para estimar el esfuerzo de desarrollo, se aplicó el modelo COCOMO en sus variantes Básica e Intermedia, clasificando el proyecto en la categoría ", False, False),
    ("Semi-desacoplado (Semi-detached)", True, False),
    (", correspondiente a proyectos de mediana complejidad con equipos de experiencia mixta.", False, False),
])

# --- DATOS Y CONTEXTO ---
add_p("Datos y contexto del proyecto:", bold=True)

make_table_teal(
    ["Parámetro", "Valor"],
    [
        ["Proyecto", "Leo-Connect (Lácteos Leo)"],
        ["Tipo de Proyecto / Modo", "Semi-desacoplado (Semi-detached)"],
        ["Equipo de Desarrollo", "4 Estudiantes Universitarios (5to Semestre – UNACH)"],
        ["Stack Tecnológico", "React 19 + TypeScript + Supabase (PostgreSQL) + Vite 8 + Tailwind CSS"],
        ["Tamaño Estimado (KLOC)", "15"],
        ["Tamaño (líneas de código)", "15,000"],
    ]
)

# --- DESGLOSE POR MÓDULO ---
add_p("Desglose del tamaño por módulo (KLOC):", bold=True)

make_table_teal(
    ["Módulo del Sistema", "Componentes Principales", "KLOC Estimadas", "% del Total"],
    [
        ["1. E-commerce (Tienda Virtual)", "Catálogo, Carrito de compras, Checkout, IVA", "3.5", "23.3%"],
        ["2. Facturación Electrónica SRI", "Estructura XML/XSD, Clave acceso 49 dígitos, RIDE PDF", "3.5", "23.3%"],
        ["3. Autenticación y Seguridad", "Validador Cédula (Módulo 10), Roles, RLS, JWT, OAuth Google", "2.5", "16.7%"],
        ["4. Gestión de Inventario", "Descuento de stock, CRUD de productos, alertas", "1.5", "10.0%"],
        ["5. Reportes y Exportación", "Panel administrativo con gráficos, Exportar CSV/PDF", "1.5", "10.0%"],
        ["6. Infraestructura Core", "Enrutador (Router), Controladores y Vistas base, Contextos React", "2.5", "16.7%"],
        ["TOTAL ESTIMADO", "Sistema Completo Leo-Connect", "15", "100%"],
    ]
)

# --- PARTE 1: COCOMO BÁSICO ---
add_p("PARTE 1: ESTIMACIÓN COCOMO BÁSICO (Condiciones Genéricas)", bold=True)

add_p("Constantes del Modelo Semi-desacoplado (Semi-detached):")
make_table_teal(
    ["Constante", "Símbolo", "Valor"],
    [
        ["Coeficiente de Esfuerzo", "a", "3"],
        ["Exponente de Esfuerzo", "b", "1.12"],
        ["Coeficiente de Tiempo", "c", "2.5"],
        ["Exponente de Tiempo", "d", "0.35"],
    ]
)

add_p("Resultados – COCOMO Básico:", bold=True)
add_formula([
    "Esfuerzo (E) = a × (KLOC)^b",
    "E = 3.0 × (15)^1.12 = 62.28 meses-persona",
    "",
    "Tiempo de Desarrollo (Tdev) = c × (E)^d",
    "Tdev = 2.5 × (62.28)^0.35 = 10.62 meses",
    "",
    "Personal Requerido = E / Tdev",
    "P = 62.28 / 10.62 = 5.87 ≈ 6 desarrolladores estándar",
])

# --- PARTE 2: COCOMO INTERMEDIO ---
add_p("PARTE 2: COCOMO INTERMEDIO (Análisis Real de Nuestro Equipo)", bold=True)

add_p("Selección de los 7 Factores de Ajuste del Esfuerzo (FAE) según la realidad del equipo:")

make_table_teal(
    ["Factor COCOMO", "Nivel Asignado", "Valor (Multiplicador)", "Justificación Real del Equipo UNACH"],
    [
        ["RELY\n(Confiabilidad del Producto)", "ALTA", "1.15", "Manejamos facturación fiscal SRI y datos financieros; un fallo en XML o cálculo causa rechazo fiscal y pérdidas al negocio."],
        ["CPLX\n(Complejidad del Producto)", "ALTO", "1.15", "Alta complejidad técnica: XML/XSD SRI, clave acceso 49 dígitos, algoritmo módulo 10 para cédula, y RIDE con código de barras."],
        ["AEXP\n(Exp. en Aplicaciones Similares)", "BAJA", "1.13", "Como estudiantes de 5to semestre, ninguno ha desarrollado previamente un e-commerce + facturación fiscal conectado a APIs/WS del SRI."],
        ["PCAP\n(Capacidad de los Analistas)", "NOMINAL", "1.00", "Somos 4 estudiantes dedicados con sólida formación universitaria, aplicamos buenas prácticas de estándar de capacidad."],
        ["DATA\n(Tamaño de Base de Datos)", "BAJO", "0.94", "El catálogo de Lácteos Leo y sus clientes no representan millones de registros; volumen de datos ligero y manejable."],
        ["LEXP\n(Exp. en Lenguaje y BD)", "ALTA", "0.95", "Dominamos React, TypeScript, PostgreSQL y JavaScript gracias a materias previas (POO, Base de Datos I y Programación Web)."],
        ["TOOL\n(Uso de Herramientas)", "ALTO", "0.91", "Usamos Git, VS Code, GitHub para trabajar en equipo, Mailing y asistencia de IA (Gemini/ChatGPT) de desarrollo."],
    ]
)

add_p("FAE Total (Producto de los 7 factores):", bold=True)
add_formula([
    "FAE = 1.15 × 1.15 × 1.13 × 1.00 × 0.94 × 0.95 × 0.91",
    "FAE = 1.2164",
    "",
    "→ +21.64% de incremento neto en el esfuerzo",
    "   por el rigor fiscal del SRI",
])

add_p("Resultados – COCOMO Intermedio:", bold=True)

make_table_teal(
    ["Métrica", "Fórmula / Cálculo", "Resultado"],
    [
        ["Esfuerzo Nominal", "E_nom = a × (KLOC)^b = 3.0 × (15)^1.12", "62.28 meses-persona"],
        ["FAE Total (7 Factores)", "1.15 × 1.15 × 1.13 × 1.00 × 0.94 × 0.95 × 0.91", "1.2164"],
        ["Esfuerzo Ajustado (E_adj)", "E_adj = E_nom × FAE = 62.28 × 1.2164", "75.63 meses-persona"],
        ["Tiempo de Desarrollo (Tdev)", "Tdev = c × (E_adj)^d = 2.5 × (75.63)^0.35", "11.86 meses"],
        ["Personal Total COCOMO", "P = E_adj / Tdev = 75.63 / 11.36", "≈ 6.66 → 7 desarrolladores"],
    ]
)

# --- PARTE 3: ANÁLISIS DE LA REALIDAD ---
add_p("PARTE 3: ANÁLISIS DE LA REALIDAD DE NUESTRO EQUIPO (4 ESTUDIANTES)", bold=True)

add_mixed([
    ("¿Por qué COCOMO calcula 7 personas (6.66) pero nosotros somos solo 4 estudiantes? ¿Cómo lo logramos?", True, True),
])

make_table_teal(
    ["Factor Compensatorio", "Explicación"],
    [
        ["Personal Real Disponible", "4 Estudiantes de 5to semestre de la UNACH."],
        ["Tiempo Extra / Dedicación", "El ritmo del proyecto en la materia de Ingeniería de Software reemplazó el ritmo laboral tradicional de 160h, trabajamos con 19 horas semanales (+1.3 veces más que un desarrollador estándar)."],
        ["Uso Intensivo de IA", "Un asistente de IA (Gemini Cloud/ChatGPT) nos permitió reutilizar plantillas XML y dividir los 5 módulos en sprints ágiles de alta velocidad."],
        ["¿Cómo logramos terminar en el semestre?", "La combinación de uso intensivo de IA + distribución inteligente de módulos + metodología ágil con sprints semanales permitió que 4 personas lograran el trabajo estimado para 7."],
    ]
)

# --- ANÁLISIS COMPARATIVO ---
add_p("ANÁLISIS COMPARATIVO FINAL: BÁSICO vs INTERMEDIO", bold=True)

make_table_teal(
    ["Métrica Calculada", "COCOMO Básico\n(Genérico)", "COCOMO Intermedio\n(Equipo Real)", "Diferencia (%) y Significado"],
    [
        ["FAE (Factor de Ajuste)", "No aplica (1.000)", "1.2164", "+25.48% — El rigor fiscal del SRI eleva el esfuerzo"],
        ["Esfuerzo (meses-persona)", "62.28", "75.63", "+21.48% — Mayor dedicación requerida"],
        ["Tiempo Teórico (meses)", "10.62", "11.36", "+7.06% — 7 días más de calendario indefinido"],
        ["Personal Teórico COCOMO", "5.9", "6.7", "+13.5% — Requiere 1 desarrollador adicional estándar"],
        ["Personal Real Nuestro Equipo", "4 Estudiantes", "4 Estudiantes", "Compensamos con agilidad, distribución de módulos, uso de IA y trabajo extra"],
    ]
)

add_heading_apa("7.2.4 Herramientas de investigación utilizadas", level=3)
add_p("Se emplearon las siguientes técnicas e instrumentos de investigación:")
for bp, nt in [
    ("Entrevista semiestructurada ", "con el propietario de Lácteos Leo para el levantamiento de requisitos."),
    ("Análisis documental ", "del sitio web existente (lacteosleo.com) para identificar sus limitaciones."),
    ("Revisión bibliográfica ", "de estándares de e-commerce, seguridad web (OWASP), normas del SRI para facturación electrónica y mejores prácticas en UX/UI."),
    ("Prototipado iterativo ", "para validar las interfaces con el usuario final antes del desarrollo."),
    ("Pruebas de caja negra ", "para validar los flujos de usuario y los procesos de pago y emisión de comprobantes."),
]:
    add_bullet(nt, bold_prefix=bp)

pb()

# ============================================================
# 7.3 ACCIONES REALIZADAS
# ============================================================
add_heading_apa("7.3 Descripción de las acciones realizadas", level=2)

add_heading_apa("7.3.1 Fase de ejecución: Análisis del problema y contexto", level=3)

add_mixed([
    ("El equipo inició el proyecto con un análisis exhaustivo del contexto digital de Lácteos Leo. El sitio web preexistente (", False, False),
    ("lacteosleo.com", False, True),
    (") representaba un modelo de presencia digital mínima: una página HTML estática construida con Bootstrap 5, que funcionaba como un dashboard interactivo para mostrar información corporativa, una sección de productos con imágenes y precios, y un botón de contacto que redirigía a WhatsApp. Este modelo, aunque funcional para dar a conocer la marca, presentaba las siguientes limitaciones críticas:", False, False),
])

for lim in [
    "Ausencia de sistema de carrito de compras y proceso de pago en línea.",
    "Imposibilidad de gestionar el inventario en tiempo real.",
    "Sin sistema de autenticación ni perfiles de usuario.",
    "Sin generación automatizada de facturas o comprobantes de venta electrónica SRI.",
    "Proceso de venta completamente manual, dependiente de la disponibilidad del propietario en WhatsApp.",
    "Sin panel administrativo para la gestión integral del negocio.",
]:
    add_bullet(lim)

add_heading_apa("7.3.2 Arquitectura del sistema desarrollado", level=3)
add_mixed([
    ("La arquitectura adoptada sigue el patrón ", False, False),
    ("SPA (Single Page Application)", True, False),
    (" con separación de responsabilidades entre el frontend y el backend:", False, False),
])
for bp, nt in [
    ("Frontend: ", "Aplicación React 19 + TypeScript compilada con Vite 8, estructurada en componentes reutilizables y páginas para cada módulo del sistema."),
    ("Backend as a Service (BaaS): ", "Supabase, que provee base de datos PostgreSQL, autenticación OAuth 2.0, almacenamiento de archivos y API REST autogenerada con Row Level Security."),
    ("Despliegue: ", "Vercel (CDN global con HTTPS automático y CI/CD desde GitHub)."),
]:
    add_bullet(nt, bold_prefix=bp)

add_heading_apa("7.3.3 Módulos desarrollados", level=3)

# Módulos
modulos = [
    ("a) Módulo de E-commerce (Tienda Virtual) – 3.5 KLOC", [
        "Catálogo dinámico con los 23 productos de Lácteos Leo organizados en cuatro categorías: Quesos, Yogures, Bebidas y Especiales.",
        "Filtrado por categoría y búsqueda en tiempo real por nombre de producto.",
        "Tarjetas de producto con imagen, nombre, descripción, precio, unidad, badge de clasificación y estado de stock.",
        "Carrito de compras persistente implementado con Context API de React con cálculo de IVA 15%.",
        "Proceso de Checkout con pasarela de pago simulada (Kushki/PayPhone) con protocolo 3D Secure 2.0 (EMVCo).",
        "Challenge de OTP por SMS o verificación biométrica y generación de código de autorización único.",
    ]),
    ("b) Módulo de Facturación Electrónica SRI – 3.5 KLOC", [
        "Generación de XML de factura electrónica con clave de acceso de 49 dígitos (fecha, tipo de comprobante, RUC del emisor, número de serie y dígito verificador módulo 11).",
        "Generación de RIDE (Representación Impresa del Documento Electrónico) en formato PDF utilizando jsPDF y jspdf-autotable.",
        "Almacenamiento del comprobante en la tabla comprobantes_sri de la base de datos.",
        "Envío del comprobante al correo del cliente mediante función Edge de Supabase.",
        "Módulo \"Mis Facturas\" donde el cliente puede consultar el historial de sus compras y descargar los comprobantes en PDF.",
    ]),
    ("c) Módulo de Autenticación y Seguridad – 2.5 KLOC", [
        "Registro e inicio de sesión con correo electrónico y contraseña, validando formato de cédula ecuatoriana con algoritmo módulo 10.",
        "Inicio de sesión y registro mediante Google OAuth 2.0.",
        "Sistema multirol con tres perfiles: Administrador (gestión total), Vendedor (panel POS) y Cliente (catálogo y compras).",
        "Políticas de Row Level Security (RLS) en Supabase para garantizar que cada rol acceda únicamente a los datos autorizados (12 políticas RLS implementadas).",
        "Persistencia de sesión mediante tokens JWT y respaldo en localStorage.",
        "Modal de autenticación integrado en el Navbar, accesible desde cualquier página.",
    ]),
    ("d) Módulo de Gestión de Inventario – 1.5 KLOC", [
        "CRUD completo (Crear, Leer, Actualizar, Eliminar) de los 23 productos del catálogo.",
        "Descuento automático del stock de los productos comprados en la base de datos.",
        "Alertas de stock bajo (productos con stock ≤ 15 unidades).",
        "Actualización de precios e imágenes en tiempo real.",
    ]),
    ("e) Módulo de Reportes y Exportación – 1.5 KLOC", [
        "Panel administrativo (Dashboard) con métricas en tiempo real: total de productos, pedidos, usuarios y ventas totales.",
        "Gráficos de barras y líneas de tendencia renderizados con Canvas API sin dependencias externas.",
        "Exportación de reportes de ventas en formato PDF y CSV.",
        "Panel de Vendedor (POS) con sistema de punto de venta para ventas presenciales con registro automático de pedidos.",
    ]),
    ("f) Infraestructura Core – 2.5 KLOC", [
        "Navbar responsiva con logo, menú de secciones, carrito con contador dinámico y accesos diferenciados según el rol.",
        "TopBar con información de contacto rápido y FloatingWhatsApp para contacto directo.",
        "Hero Section con video de fondo a pantalla completa, animaciones AOS y llamada a la acción.",
        "Sección Historia de la empresa y Sección Contacto con formulario.",
        "VideoShowcase: galería multimedia de videos demostrativos de los productos.",
        "Enrutador con React Router DOM v7, contextos (AuthContext, CartContext) y utilidades (sriGenerator).",
    ]),
]

for title, items in modulos:
    add_p(title, bold=True)
    for item in items:
        add_bullet(item)

add_heading_apa("7.3.4 Fase de socialización y reflexión", level=3)

add_p("Durante la fase de socialización, el equipo presentó los avances del proyecto en sesiones semanales de seguimiento con el docente. La reflexión crítica del proceso permitió identificar las siguientes lecciones aprendidas:")

for l in [
    "La integración de herramientas de IA generativa (Google Gemini y ChatGPT) como par de programación aceleró significativamente el desarrollo de componentes complejos como el generador de XML del SRI con clave de acceso de 49 dígitos y el módulo de 3D Secure, que de otra manera habrían requerido semanas adicionales de investigación.",
    "El uso de Supabase como BaaS eliminó la necesidad de desarrollar un backend completo desde cero, permitiendo al equipo de 4 estudiantes concentrarse en la lógica de negocio y la experiencia de usuario.",
    "La gestión de roles y permisos con 12 políticas RLS demostró ser un enfoque robusto y seguro para controlar el acceso diferenciado a los datos.",
    "La adopción de TypeScript desde el inicio del proyecto redujo notablemente los errores en tiempo de ejecución gracias al tipado estático.",
    "El modelo COCOMO Intermedio validó que el FAE de 1.2164 incrementa el esfuerzo un +21.64% respecto al cálculo genérico, confirmando que 4 estudiantes con IA pueden lograr el trabajo estimado para 7 programadores convencionales.",
]:
    add_bullet(l)

pb()

# ============================================================
# 7.4 RESULTADOS
# ============================================================
add_heading_apa("7.4 Resultados", level=2)

add_heading_apa("7.4.1 Plataforma web funcional desplegada en producción", level=3)

add_mixed([
    ("El principal resultado del proyecto es la plataforma ", False, False),
    ("LEO-CONNECT", True, False),
    (", una aplicación web de e-commerce con facturación electrónica SRI completamente funcional que ha sido desplegada en producción. A continuación se presentan los resultados por módulo:", False, False),
])

pruebas = [
    ("Prueba 1: Página Principal y Catálogo", "Se verificó la carga correcta del Hero Section con video de fondo, la navegación entre secciones (Inicio, Nosotros, Catálogo, Contacto) y la visualización del catálogo de 23 productos con filtros por categoría activos. Resultado: APROBADO.",
     ["Página Principal – Hero Section con video de Lácteos Leo", "Catálogo de Productos con filtros por categoría"]),
    ("Prueba 2: Proceso de Registro e Inicio de Sesión", "Se validó el registro de un nuevo usuario con correo/contraseña y cédula ecuatoriana (validación módulo 10), así como el inicio de sesión con Google OAuth. Se verificó la asignación automática de roles. Resultado: APROBADO.",
     ["Modal de Autenticación – Registro e Inicio de Sesión con Google OAuth"]),
    ("Prueba 3: Carrito de Compras y Checkout con 3D Secure", "Se probó el flujo completo de compra: agregar productos al carrito, verificar validación de stock, proceder al checkout con simulación 3D Secure 2.0 (OTP de 6 dígitos). Resultado: APROBADO.",
     ["Carrito de Compras con productos y totales", "Proceso de Pago con 3D Secure – Ingreso de OTP"]),
    ("Prueba 4: Generación de Comprobante Electrónico SRI", "Se verificó la generación correcta del XML de factura electrónica con clave de acceso de 49 dígitos, el PDF RIDE y el almacenamiento en la base de datos. Resultado: APROBADO.",
     ["Comprobante Electrónico PDF (RIDE) generado", "Módulo Mis Facturas con historial de compras"]),
    ("Prueba 5: Panel de Administración", "Se validó el acceso restringido al panel de administración solo para el rol administrador. Se probaron Dashboard, CRUD de productos, gestión de usuarios y reportes. Resultado: APROBADO.",
     ["Dashboard de Administración con métricas y gráficos", "Gestión de Productos – CRUD de inventario"]),
    ("Prueba 6: Panel de Vendedor (POS)", "Se probó el flujo completo del punto de venta: búsqueda de productos, carrito de venta, validación de stock y registro automático del pedido. Resultado: APROBADO.",
     ["Panel de Vendedor – Sistema POS"]),
]

for titulo, desc, caps in pruebas:
    add_p(titulo, bold=True)
    add_p(desc)
    for cap in caps:
        add_placeholder(f"CAPTURA DE PANTALLA: {cap}", f"Insertar screenshot: {cap}")

add_heading_apa("7.4.2 Base de datos relacional implementada", level=3)
add_p("Se implementó un esquema de base de datos relacional en PostgreSQL (Supabase) con las siguientes tablas:")
make_table(
    ["Tabla", "Descripción", "Registros iniciales"],
    [
        ["roles", "Roles del sistema: administrador, vendedor, cliente", "3"],
        ["usuarios", "Usuarios registrados con rol, cédula y datos de contacto", "Variable"],
        ["productos", "Catálogo de productos con stock, precio e imágenes", "23"],
        ["pedidos", "Pedidos realizados con método de pago y código de autorización", "Variable"],
        ["detalle_pedidos", "Líneas de detalle de cada pedido", "Variable"],
        ["comprobantes_sri", "Comprobantes electrónicos con clave de acceso y XML", "Variable"],
    ]
)

add_heading_apa("7.4.3 Catálogo de productos implementado", level=3)
make_table(
    ["Categoría", "Productos", "Precio rango"],
    [
        ["Quesos", "Queso Maduro, Tierno, Cuadrado, Familiar, Redondo, Mozzarella Pequeño/Entero/Bolita, Cheddar, Bloque Mozzarella 2.5kg", "$0.80 – $12.00"],
        ["Yogures", "Funda Yogurt, Yogurt c/ Cereal, ½Lt, 1L, 2L, 4L (Galón)", "$0.50 – $3.25"],
        ["Bebidas", "Naranjadas, Colas, Bolos, Bebas Refrescantes", "$1.00 – $1.50"],
        ["Especiales", "Gelatina, Bolita de Mantequilla, Queso Manaba", "$1.00 – $2.50"],
    ]
)

add_heading_apa("7.4.4 Métricas del proyecto", level=3)
make_table(
    ["Métrica", "Valor"],
    [
        ["Total de archivos fuente (TSX/TS/CSS)", "35+ archivos"],
        ["Líneas de código estimadas (LOC)", "~15,000 LOC (15 KLOC)"],
        ["Componentes React implementados", "18 componentes"],
        ["Páginas de la aplicación", "8 páginas / vistas"],
        ["Tablas en base de datos", "6 tablas"],
        ["Productos del catálogo", "23 productos"],
        ["Roles de usuario", "3 roles"],
        ["Políticas RLS en Supabase", "12 políticas"],
        ["Dependencias npm utilizadas", "10 dependencias"],
        ["Plataforma de despliegue", "Vercel (HTTPS, CDN global)"],
        ["Esfuerzo COCOMO Básico", "62.28 personas-mes"],
        ["Esfuerzo COCOMO Intermedio (ajustado)", "75.63 personas-mes"],
        ["FAE Total", "1.2164 (+21.64%)"],
        ["Personal COCOMO teórico", "≈ 7 programadores"],
        ["Equipo de desarrollo real", "4 estudiantes UNACH + IA"],
    ]
)

pb()

# ============================================================
# 7.5 BIBLIOGRAFÍA
# ============================================================
add_heading_apa("7.5 Bibliografía", level=2)

refs = [
    "Boehm, B. W. (1981). Software engineering economics. Prentice-Hall.",
    "Chacon, S., & Straub, B. (2014). Pro Git (2.ª ed.). Apress. https://git-scm.com/book/en/v2",
    "Google. (2025). Gemini API documentation. Google AI for Developers. https://ai.google.dev/docs",
    "Kushki. (2025). Documentación técnica de la pasarela de pagos Kushki para Ecuador. https://docs.kushkipagos.com",
    "Lácteos Leo. (2024). Sitio web oficial de Lácteos Leo. https://lacteosleo.com",
    "Meta Open Source. (2025). React documentation (versión 19). https://react.dev",
    "Mozilla Developer Network. (2025). Web technology for developers: HTML, CSS, JavaScript. https://developer.mozilla.org",
    "OWASP Foundation. (2021). OWASP Top 10: 2021 – The ten most critical web application security risks. https://owasp.org/Top10",
    "Pressman, R. S., & Maxim, B. R. (2020). Ingeniería del software: Un enfoque práctico (8.ª ed.). McGraw-Hill.",
    "Servicio de Rentas Internas (SRI). (2024). Ficha técnica de comprobantes electrónicos versión 2.21. Gobierno del Ecuador. https://www.sri.gob.ec/facturacion-electronica",
    "Sommerville, I. (2016). Ingeniería de software (10.ª ed.). Pearson.",
    "Supabase, Inc. (2025). Supabase documentation: Open source Firebase alternative. https://supabase.com/docs",
    "Vercel, Inc. (2025). Vercel platform documentation. https://vercel.com/docs",
    "Vite. (2025). Vite: Next generation frontend tooling (versión 8). https://vite.dev",
    "W3C. (2023). Accessibility guidelines (WCAG) 2.2. World Wide Web Consortium. https://www.w3.org/TR/WCAG22",
]
for r in refs:
    add_ref(r)

pb()

# ============================================================
# 8. ANEXOS
# ============================================================
add_heading_apa("8. ANEXOS (EVIDENCIAS)", level=1)

add_heading_apa("Anexo A – Diagrama de Casos de Uso", level=2)
add_placeholder("DIAGRAMA DE CASOS DE USO", "Insertar el diagrama de casos de uso.\nActores: Cliente, Vendedor, Administrador, Sistema SRI, Pasarela de Pago\nUC-01 a UC-12: Registro, Catálogo, Carrito, Checkout 3DS, Comprobante SRI, Facturas, Inventario, Dashboard, Usuarios, Reportes, POS, WhatsApp")

add_heading_apa("Anexo B – Diagrama de Clases", level=2)
add_placeholder("DIAGRAMA DE CLASES UML", "Insertar el diagrama de clases.\nClases: Usuario, Rol, Producto, Pedido, DetallePedido, ComprobanteSRI, CartItem\nRelaciones: Usuario *–1 Rol | Pedido *–1 Usuario | DetallePedido *–1 Pedido | ComprobanteSRI 1–1 Pedido")

add_heading_apa("Anexo C – Diagrama Entidad-Relación (ERD)", level=2)
add_placeholder("DIAGRAMA ENTIDAD-RELACIÓN", "Insertar el diagrama ERD.\nEntidades: roles, usuarios, productos, pedidos, detalle_pedidos, comprobantes_sri\nRestricciones CHECK en precio, stock y cantidad. RLS habilitado en todas las tablas.")

add_heading_apa("Anexo D – Diagrama de Arquitectura del Sistema", level=2)
add_placeholder("DIAGRAMA DE ARQUITECTURA", "Insertar el diagrama de arquitectura.\nCapas: Usuario → Vercel CDN → Frontend (React 19 + TS + Vite 8) → Supabase Backend (PostgreSQL, Auth, Edge Functions) → Servicios Externos (Google OAuth, Pasarela, SRI)")

add_heading_apa("Anexo E – Diagrama de Secuencia: Flujo de Compra y Facturación", level=2)
add_placeholder("DIAGRAMA DE SECUENCIA", "Insertar el diagrama de secuencia.\n1. Cliente selecciona producto → 2. addToCart() → 3. Checkout\n4. Verificar autenticación → 5-6. Datos y método pago\n7-11. Pasarela 3D Secure (OTP) → 12. crearPedidoCompleto()\n13-14. generateXML() + generateRIDE() → 15. Email → 16. Éxito")

add_heading_apa("Anexo F – Diagrama de Actividades: Autenticación Multirol", level=2)
add_placeholder("DIAGRAMA DE ACTIVIDADES", "Insertar el diagrama de actividades.\nINICIO → ¿Tiene cuenta? → Login / Registro → Supabase Auth → fetchUserRole()\n→ ¿Admin? → /admin | ¿Vendedor? → /vendedor | ¿Cliente? → / → FIN")

add_heading_apa("Anexo G – Estimación COCOMO Completa", level=2)
add_p("Las tablas completas de la estimación COCOMO Básico e Intermedio se presentaron en la sección 7.2.3 de este informe. A continuación se incluye la imagen original de la hoja de cálculo COCOMO elaborada por el equipo:")
add_placeholder("IMAGEN: Hoja de Cálculo COCOMO Original del Equipo", "Insertar aquí la imagen/captura de la hoja Excel COCOMO con los datos del proyecto Leo-Connect (la tabla verde-azulada con Parte 1, Parte 2, Parte 3 y Análisis Comparativo)")

add_heading_apa("Anexo H – Evidencias del Trabajo en Equipo", level=2)

for title, desc in [
    ("Repositorio Git/GitHub", "Insertar captura del repositorio con historial de commits y enlace público al código fuente"),
    ("Reuniones del equipo", "Insertar fotos o capturas de las sesiones de planificación y seguimiento semanal"),
    ("Plataforma Supabase en producción", "Insertar screenshot del dashboard de Supabase con las 6 tablas del proyecto"),
    ("Despliegue en Vercel", "Insertar screenshot del dashboard de Vercel mostrando el proyecto desplegado"),
    ("Mockups / Prototipos de interfaces", "Insertar imágenes de los wireframes y prototipos de alta fidelidad elaborados"),
    ("Sitio web original (lacteosleo.com)", "Insertar screenshot del sitio preexistente para comparar el ANTES y el DESPUÉS del proyecto"),
]:
    add_placeholder(f"EVIDENCIA: {title}", desc)

# ============================================================
# GUARDAR
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "INFORME_LEO_CONNECT.docx")
doc.save(output_path)
print(f"\n{'='*60}")
print(f"  DOCUMENTO WORD GENERADO EXITOSAMENTE")
print(f"  Ubicación: {output_path}")
print(f"{'='*60}\n")
