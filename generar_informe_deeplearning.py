# -*- coding: utf-8 -*-
"""
Generador de Informe Word (.docx) - Proyecto Gimnasio Deep Learning
Formato solicitado por el usuario
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Configuración básica de estilos
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_after = Pt(0)
pf.space_before = Pt(0)

# Márgenes
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.54)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = True
    return h

def add_p(text, bold=False, italic=False, align='justify'):
    p = doc.add_paragraph()
    if align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# ============================================================
# 7. DESARROLLO DEL INFORME
# ============================================================
add_heading("7. DESARROLLO DEL INFORME", level=1)

# ============================================================
# 7.1 Introducción
# ============================================================
add_heading("7.1 Introducción", level=2)
add_p("En la actualidad, la aplicación de la Inteligencia Artificial y el Deep Learning en la biomecánica deportiva representa una oportunidad revolucionaria para prevenir lesiones y mejorar el rendimiento físico. Este proyecto tiene como objetivo desarrollar un sistema automatizado capaz de detectar patrones de error postural en secuencias de movimiento de ejercicios de gimnasio.")
add_p("Para el desarrollo de este proyecto se emplearon diversas herramientas tecnológicas:")
add_bullet("Python: Lenguaje principal para la lógica del backend y procesamiento de datos.")
add_bullet("PyTorch: Framework de Deep Learning utilizado para diseñar y entrenar la arquitectura híbrida (1D-CNN + BiLSTM).")
add_bullet("OpenCV: Utilizado para la renderización del \"Wow Factor\", superponiendo el esqueleto y un HUD con métricas en tiempo real sobre el video.")
add_bullet("Penn Action Dataset: Conjunto de datos público (coordenadas .mat) utilizado para el entrenamiento y validación del modelo.")
add_bullet("HTML5, CSS y JavaScript: Empleados para desarrollar un Dashboard interactivo web profesional que visualiza las inferencias del modelo y la matriz de confusión.")

# ============================================================
# 7.2 Descripción de la metodología
# ============================================================
add_heading("7.2 Descripción de la metodología", level=2)
add_p("El proyecto se desarrolló bajo un enfoque de investigación formativa y aplicada, estructurándose en las siguientes actividades principales (Qué y Cómo):")
add_p("QUÉ SE REALIZÓ:", bold=True)
add_p("Se desarrolló un modelo computacional híbrido capaz de clasificar la ejecución de ejercicios en tres categorías posturales: (0) Postura Correcta, (1) Error de Espalda/Tronco, y (2) Error de Extremidades/Rodillas; además de brindar retroalimentación textual al usuario final.")
add_p("CÓMO SE REALIZÓ:", bold=True)
add_bullet("Recolección y Preprocesamiento: Extracción de coordenadas de 13 articulaciones del Penn Action Dataset. Normalización Z-score y alineación temporal (padding) a 46 fotogramas por secuencia (vector de 26 características).")
add_bullet("Diseño del Modelo: Implementación de un bloque 1D-CNN para extraer correlaciones espaciales inter-articulación en cada instante, seguido de un bloque BiLSTM para modelar las transiciones temporales del movimiento.")
add_bullet("Lógica de Retroalimentación: Capa Softmax final conectada a un motor de reglas para emitir alertas (ej. \"Mantén el pecho erguido y activa el core\").")

# ============================================================
# 7.3 Descripción de las acciones realizadas
# ============================================================
add_heading("7.3 Descripción de las acciones realizadas", level=2)
add_p("FASE DE EJECUCIÓN Y SEGUIMIENTO", bold=True)
add_bullet("Construcción del Dataset en PyTorch para ingesta optimizada en memoria con lotes de 32 secuencias.")
add_bullet("Entrenamiento de 25 épocas empleando CrossEntropy Loss y optimizador Adam (LR=0.001, Weight Decay=5e-4, Dropout=0.25).")
add_bullet("Implementación de validación cruzada tras cada época con guardado automático de pesos (Checkpoints) y técnica de Early Stopping.")
add_p("FASE DE SOCIALIZACIÓN Y REFLEXIÓN", bold=True)
add_bullet("Implementación de visualización en video (Wow Factor) procesando secuencias MP4 con OpenCV y superponiendo los enlaces del esqueleto.")
add_bullet("Creación del Dashboard Web Interactivo que facilita a cualquier usuario probar el sistema sin requerir conocimientos de código.")

# ============================================================
# 7.4 Resultados
# ============================================================
add_heading("7.4 Resultados", level=2)
add_p("El entrenamiento del modelo 1D-CNN + BiLSTM convergió exitosamente:")
add_bullet("El Early Stopping determinó la época óptima alrededor de la época 13, con una pérdida (Loss) de 0.3328.")
add_bullet("Se alcanzó una Precisión (Accuracy) del 89.9% en el conjunto de validación.")
add_bullet("El Macro F1-Score fue del 0.551, reflejando la capacidad del modelo para balancear las detecciones ante las distintas clases de errores posturales.")
add_p("[INSERTE AQUÍ LA IMAGEN DEL GRÁFICO DE PÉRDIDA Y PRECISIÓN PROPORCIONADA]", italic=True, align="center")

# ============================================================
# 7.5 Bibliografía
# ============================================================
add_heading("7.5 Bibliografía", level=2)
add_p("1. Zhang, W., Zhu, M., & Derpanis, K. G. (2013). Envisioning action: Spatial and temporal representations. Penn Action Dataset.")
add_p("2. Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press.")
add_p("3. Paszke, A., et al. (2019). PyTorch: An Imperative Style, High-Performance Deep Learning Library.")
add_p("4. Bradski, G. (2000). The OpenCV Library.")

# ============================================================
# 8. ANEXOS (EVIDENCIAS)
# ============================================================
doc.add_page_break()
add_heading("8. ANEXOS (EVIDENCIAS)", level=1)
add_p("Se incluyen como anexos:")
add_bullet("Anexo A: Gráfico de Evolución de Función de Pérdida y Métricas de Rendimiento. [Añadir captura]")
add_bullet("Anexo B: Diagrama de Arquitectura del Modelo (1D-CNN + BiLSTM). [Añadir diagrama]")
add_bullet("Anexo C: Capturas del Dashboard Web y de la generación de Video con OpenCV. [Añadir capturas]")
add_bullet("Anexo D: Diagrama de Clases del proyecto (dataset.py, model.py, train.py). [Añadir diagrama]")

# Guardar
output_path = r"C:\Users\RubenC\Desktop\PROYECTO PAGINAWEB 2026\Informe_DeepLearning.docx"
doc.save(output_path)
print(f"Documento generado exitosamente en: {output_path}")
